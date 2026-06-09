#!/usr/bin/env python3
"""ingest/documents.py — docx / pdf 歌词本 → 文本。

歌词本除了图片，也常以 docx / pdf 提交，内含歌词 + 歌曲源信息(staff/发行等)。
本模块统一抽取其文字：

- DOCX：python-docx 解析段落与表格文字
- PDF ：先抽文本层（PyMuPDF）；若某页几乎无文字（扫描件），把该页渲染成图片
        回退到 ingest.ocr 做视觉识别

用法：
    python -m ingest.documents <booklet.pdf|booklet.docx> [--out text.txt]
    python -m ingest.documents --dir <投递目录>

依赖：PyMuPDF(fitz)、python-docx（见 ingest/requirements.txt）。缺依赖时报清晰错误。
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

if __package__ in (None, ""):
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from ingest import ocr as ocr_mod  # type: ignore
else:
    from . import ocr as ocr_mod

DOC_EXTS = {".pdf", ".docx"}
# 一页文本少于该字符数视为「可能是扫描件」，回退 OCR
_SCANNED_PAGE_THRESHOLD = 8


def find_documents(directory: Path) -> list[Path]:
    return sorted(
        p for p in directory.rglob("*") if p.is_file() and p.suffix.lower() in DOC_EXTS
    )


def extract_docx(path: Path) -> str:
    try:
        import docx  # type: ignore  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "未安装 python-docx，请先安装 ingest/requirements.txt"
        ) from e
    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    # 表格里常放分轨 staff
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def extract_pdf(path: Path) -> str:
    """抽 PDF 文本；扫描页回退到 OCR。"""
    try:
        import fitz  # type: ignore  # PyMuPDF
    except ImportError as e:
        raise RuntimeError("未安装 PyMuPDF，请先安装 ingest/requirements.txt") from e

    out_pages: list[str] = []
    with fitz.open(str(path)) as doc:
        for i, page in enumerate(doc):
            text = (page.get_text() or "").strip()
            if len(text) >= _SCANNED_PAGE_THRESHOLD:
                out_pages.append(text)
                continue
            # 扫描页：渲染成图 → OCR
            print(f"○ PDF 第 {i+1} 页文本稀少，回退 OCR", file=sys.stderr, flush=True)
            try:
                pix = page.get_pixmap(dpi=200)
                tmp = path.with_name(f".{path.stem}_p{i+1}.png")
                pix.save(str(tmp))
                try:
                    ocr_text = ocr_mod.ocr_image(tmp)
                finally:
                    tmp.unlink(missing_ok=True)
                if ocr_text:
                    out_pages.append(ocr_text)
            except Exception as e:  # noqa: BLE001
                print(f"⚠️  PDF 第 {i+1} 页 OCR 失败: {e}", file=sys.stderr, flush=True)
    return "\n\n".join(out_pages).strip()


def extract(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".pdf":
        return extract_pdf(path)
    raise ValueError(f"不支持的文档类型: {path.name}")


def run(docs: list[Path]) -> dict[str, str]:
    out: dict[str, str] = {}
    for d in docs:
        try:
            text = extract(d)
        except Exception as e:  # noqa: BLE001 — 单文件失败不中断
            print(f"⚠️  文档抽取失败 {d.name}: {e}", file=sys.stderr, flush=True)
            continue
        if text:
            out[d.name] = text
            print(f"✓ 文档 {d.name} ({len(text)} 字)", file=sys.stderr, flush=True)
    return out


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="docx/pdf 歌词本抽取")
    ap.add_argument("docs", nargs="*", help="文档路径")
    ap.add_argument("--dir", help="抽取该目录下所有 docx/pdf")
    ap.add_argument("--out", help="写入文本文件（默认 stdout）")
    args = ap.parse_args(argv)

    docs: list[Path] = [Path(p) for p in args.docs]
    if args.dir:
        docs += find_documents(Path(args.dir))
    docs = [p for p in docs if p.is_file()]
    if not docs:
        print("没有可抽取的文档", file=sys.stderr)
        return 0

    result = run(docs)
    text_out = "\n\n".join(f"# === {name} ===\n{txt}" for name, txt in result.items())
    if args.out:
        Path(args.out).write_text(text_out + "\n", encoding="utf-8")
        print(f"✓ 写入 {args.out}", file=sys.stderr)
    else:
        print(text_out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
